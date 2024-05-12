import os
import re

import fitz
import numpy as np
import requests
import textract
from decouple import config
from django.core.exceptions import ObjectDoesNotExist
from docx import Document
from sentence_transformers import SentenceTransformer, util

from Core.models import UserResume


class ParsingFunctions:
    def __init__(self):
        self.HF_TOKEN = config('HF_TOKEN')
        self.API_URL = "https://api-inference.huggingface.co/models/mistralai/Mistral-7B-Instruct-v0.2"
        self.headers = {"Authorization": f"Bearer {self.HF_TOKEN}"}
        self.model = SentenceTransformer('all-MiniLM-L6-v2')

    def query(self, payload):
        try:
            response = requests.post(self.API_URL, headers=self.headers, json=payload)
            response.raise_for_status()
            return response.json()
        except requests.HTTPError as e:
            print(f"HTTP error occurred: {e}")
        except requests.RequestException as e:
            print(f"Request exception occurred: {e}")
        except Exception as e:
            print(f"Some other error occurred: {e}")
        return []

    def gen_query(self, resume_text):
        payload = {
            "inputs":
                f"""
                Here is a resume detailing various technical experiences and skills. Please analyze the resume and extract the most important technical skills. Focus on identifying specific technologies or tools that are represented as one-word items typical in the tech industry, such as programming languages or widely recognized technology terms.:
                [{resume_text}]
                Please list the top skills demonstrated by this individual as single words that are commonly used as standalone skills in the tech industry.
                """,
            "parameters": {"return_full_text": False, "num_results": 10},
        }
        return self.query(payload)

    def post_processing(self, output):
        text = output[0]['generated_text'] if output else ''
        skills_list = [re.sub(r'^\s*\d+\.\s*', '', line).strip() for line in text.split('\n') if line.strip()]
        print("Extracted Skills:", skills_list)
        with open("Data/SkillsDataSet", 'r') as file:
            skill_dataset = [line.strip() for line in file]
        dataset_embeddings = self.model.encode(skill_dataset)
        extracted_embeddings = self.model.encode(skills_list)
        high_similarity_skills = []
        for skill, embedding in zip(skills_list, extracted_embeddings):
            similarities = util.pytorch_cos_sim(embedding, dataset_embeddings)[0]
            top_indices = np.where(similarities > 0.75)[0]
            if len(top_indices) > 0:
                highest_index = top_indices[np.argmax(similarities[top_indices])]
                high_similarity_skills.append(skill_dataset[highest_index])
        return high_similarity_skills

    def normalize_skills(self, skills):
        with open("Data/SkillsDataSet", 'r') as file:
            skill_dataset = [line.strip() for line in file]
        extracted_embeddings = self.model.encode(skills)
        dataset_embeddings = self.model.encode(skill_dataset)
        high_similarity_skills = []
        for skill, embedding in zip(skills, extracted_embeddings):
            similarities = util.pytorch_cos_sim(embedding, dataset_embeddings)[0]
            top_indices = np.where(similarities > 0.75)[0]
            if len(top_indices) > 0:
                highest_index = top_indices[np.argmax(similarities[top_indices])]
                high_similarity_skills.append(skill_dataset[highest_index])
            else:
                high_similarity_skills.append(skill)  # Append original skill if no match found
        return high_similarity_skills


class ResumeParsing:
    def __init__(self, request):
        self.request = request

    def extract_text_from_pdf(self):
        try:
            resume = self.request.user.resumes.order_by('-uploaded_at').first()
            if not resume.resume:
                return None
            text = ""
            with fitz.open(resume.resume.path) as doc:
                for page in doc:
                    text += page.get_text()
            return text
        except ObjectDoesNotExist:
            print("Resume file does not exist for the user.")
            return None
        except Exception as e:
            print(f"An error occurred while extracting text from PDF: {e}")
            return None

    def extract_text_from_docx(self):
        try:
            resume = self.request.user.resumes.order_by('-uploaded_at').first()
            if not resume.resume:
                return None
            if resume.resume.path.endswith('.docx'):
                doc = Document(resume.resume.path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                return text
            return None
        except ObjectDoesNotExist:
            print("Resume file does not exist for the user.")
            return None
        except Exception as e:
            print(f"An error occurred while extracting text from DOCX: {e}")
            return None

    def extract_text_from_doc(self):
        try:
            resume = self.request.user.resumes.order_by('-uploaded_at').first()
            if not resume.resume:
                return None
            if resume.resume.path.endswith('.doc'):
                text = textract.process(resume.resume.path).decode('utf-8')
                return text
            return None
        except ObjectDoesNotExist:
            print("Resume file does not exist for the user.")
            return None
        except Exception as e:
            print(f"An error occurred while extracting text from DOC: {e}")
            return None

    def extract_text_from_txt(self):
        try:
            resume = self.request.user.resumes.order_by('-uploaded_at').first()
            if not resume.resume:
                return None

            if resume.resume.path.endswith('.txt'):
                with open(resume.resume.path, 'r') as f:
                    text = f.read()
                return text.strip()

            return None
        except ObjectDoesNotExist:
            print("Resume file does not exist for the user.")
            return None
        except Exception as e:
            print(f"An error occurred while extracting text from TXT: {e}")
            return None


class NewResumeParsing:
    def __init__(self, resume_file):
        self.resume_file = resume_file

    def extract_text(self):
        if not os.path.exists(self.resume_file.path):
            print("Resume file does not exist.")
            return None

        try:
            if self.resume_file.name.endswith('.pdf'):
                return self.extract_text_from_pdf()
            elif self.resume_file.name.endswith('.docx'):
                return self.extract_text_from_docx()
            elif self.resume_file.name.endswith('.doc'):
                return self.extract_text_from_doc()
            elif self.resume_file.name.endswith('.txt'):
                return self.extract_text_from_txt()
        except Exception as e:
            print(f"An error occurred while extracting text: {e}")
        return None

    def extract_text_from_pdf(self):
        text = ""
        with fitz.open(self.resume_file.path) as doc:
            for page in doc:
                text += page.get_text()
        return text

    def extract_text_from_docx(self):
        doc = Document(self.resume_file.path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    def extract_text_from_doc(self):
        return textract.process(self.resume_file.path).decode('utf-8')

    def extract_text_from_txt(self):
        with open(self.resume_file.path, 'r') as f:
            return f.read().strip()
